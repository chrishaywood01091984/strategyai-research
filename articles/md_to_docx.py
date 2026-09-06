"""Convert the Q3 research markdown into a styled Word document.

Not a plain dump: headings, real Word tables with a header row, bold runs inside
paragraphs, and the brand navy/gold. The point is that Chris can open it, read it
and hand it to someone without it looking like a text file that lost a fight.
"""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"C:/Users/ChrisHaywood/Code/strategyai-research/articles/q3-2026-mandate-formation.md"
OUT = r"C:/Users/ChrisHaywood/Code/strategyai-research/articles/The Mandate Gap - StrategyAI Research Q3 2026.docx"

NAVY = RGBColor(0x0D, 0x1B, 0x2A)
GOLD = RGBColor(0x9A, 0x7B, 0x10)   # darker gold, readable on white paper
GREY = RGBColor(0x55, 0x63, 0x7A)


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def add_runs(par, text):
    """Render **bold** and *italic* inline, and strip stray markdown."""
    for part in re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = par.add_run(part[2:-2]); r.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            r = par.add_run(part[1:-1]); r.italic = True
        else:
            par.add_run(part)


def main():
    md = open(SRC, encoding='utf-8').read()
    doc = Document()

    # Base style
    st = doc.styles['Normal']
    st.font.name = 'Georgia'
    st.font.size = Pt(10.5)
    st.paragraph_format.space_after = Pt(8)
    st.paragraph_format.line_spacing = 1.25

    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(1.0)
        s.top_margin = s.bottom_margin = Inches(0.9)

    lines = md.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Tables: a header row, a separator, then body rows.
        if line.startswith('|') and i + 1 < len(lines) and set(lines[i+1].replace('|', '').strip()) <= set('-: '):
            header = [c.strip() for c in line.strip('|').split('|')]
            aligns = [c.strip() for c in lines[i+1].strip('|').split('|')]
            i += 2
            rows = []
            while i < len(lines) and lines[i].startswith('|'):
                rows.append([c.strip() for c in lines[i].strip('|').split('|')])
                i += 1
            t = doc.add_table(rows=1, cols=len(header))
            t.style = 'Table Grid'
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            for j, h in enumerate(header):
                cell = t.rows[0].cells[j]
                cell.text = ''
                p = cell.paragraphs[0]
                r = p.add_run(re.sub(r'\*\*', '', h))
                r.bold = True; r.font.size = Pt(8.5); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.name = 'Segoe UI'
                shade(cell, '0D1B2A')
            for row in rows:
                cells = t.add_row().cells
                for j, val in enumerate(row[:len(header)]):
                    cells[j].text = ''
                    p = cells[j].paragraphs[0]
                    if aligns[j].endswith(':'):
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    bold = '**' in val
                    r = p.add_run(re.sub(r'\*\*', '', val))
                    r.font.size = Pt(9); r.font.name = 'Segoe UI'
                    if bold:
                        r.bold = True; r.font.color.rgb = GOLD
            doc.add_paragraph()
            continue

        if line.startswith('# '):
            p = doc.add_paragraph()
            r = p.add_run(line[2:]); r.bold = True; r.font.size = Pt(30)
            r.font.color.rgb = NAVY; r.font.name = 'Georgia'
            p.paragraph_format.space_after = Pt(4)
        elif line.startswith('### '):
            p = doc.add_paragraph()
            r = p.add_run(line[4:]); r.font.size = Pt(13); r.italic = True
            r.font.color.rgb = GREY; r.font.name = 'Georgia'
            p.paragraph_format.space_before = Pt(10)
        elif line.startswith('## '):
            p = doc.add_paragraph()
            r = p.add_run(line[3:]); r.bold = True; r.font.size = Pt(15)
            r.font.color.rgb = NAVY; r.font.name = 'Georgia'
            p.paragraph_format.space_before = Pt(20)
            p.paragraph_format.space_after = Pt(8)
        elif line.strip() == '---':
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pbdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:color'), 'C9A84C')
            pbdr.append(bottom); pPr.append(pbdr)
        elif line.startswith('**StrategyAI Research'):
            p = doc.add_paragraph()
            r = p.add_run(re.sub(r'\*\*', '', line))
            r.font.size = Pt(9); r.font.color.rgb = GOLD; r.font.name = 'Segoe UI'
            r.bold = True
        elif line.startswith('*') and line.endswith('*') and len(line) > 40 and not line.startswith('**'):
            p = doc.add_paragraph()
            r = p.add_run(line.strip('*'))
            r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GREY
        elif line.strip():
            # Join wrapped source lines into one paragraph.
            buf = [line]
            while i + 1 < len(lines) and lines[i+1].strip() and not lines[i+1].startswith(('#', '|', '---', '*Strategy')):
                i += 1
                buf.append(lines[i].strip())
            add_runs(doc.add_paragraph(), ' '.join(buf))
        i += 1

    doc.save(OUT)
    print("wrote:", OUT)


if __name__ == '__main__':
    main()
